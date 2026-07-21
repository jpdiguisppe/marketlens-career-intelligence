import os
from collections.abc import Generator
from io import BytesIO

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

os.environ["ADMIN_API_KEY"] = "test-admin-key"
os.environ["AUTH_DEV_MODE"] = "true"
os.environ["AUTH_DEV_BEARER_TOKEN"] = "test-user-token"
os.environ["AUTH_DEV_USER_ID"] = "test-clerk-user-1"

from app.database import Base, get_db
from app.main import _rate_limit_buckets, app

TEST_DATABASE_URL = "sqlite://"
ADMIN_HEADERS = {"X-Admin-API-Key": "test-admin-key"}
AUTH_HEADERS = {"Authorization": "Bearer test-user-token"}

engine = create_engine(
    TEST_DATABASE_URL,
    connect_args={"check_same_thread": False},
    poolclass=StaticPool,
)
TestingSessionLocal = sessionmaker(bind=engine)


def override_get_db() -> Generator[Session, None, None]:
    db = TestingSessionLocal()
    try:
        yield db
    finally:
        db.close()


app.dependency_overrides[get_db] = override_get_db
client = TestClient(app)


@pytest.fixture(autouse=True)
def reset_database() -> Generator[None, None, None]:
    Base.metadata.drop_all(bind=engine)
    Base.metadata.create_all(bind=engine)
    _rate_limit_buckets.clear()
    yield
    Base.metadata.drop_all(bind=engine)
    _rate_limit_buckets.clear()


def _sample_job_posting_payload() -> dict[str, str]:
    return {
        "company": "Comcast",
        "title": "Backend Software Engineer",
        "location": "Philadelphia, PA",
        "role_category": "Backend SWE",
        "experience_level": "Entry-Level",
        "description": "Build REST APIs with Python, SQL, Docker, Git, and Agile practices.",
    }


def _sample_docx_resume_bytes() -> bytes:
    document = Document()
    document.add_paragraph("JP Candidate")
    document.add_paragraph("Built Python FastAPI services with Docker and SQL.")
    document.add_paragraph("Used Git for version control.")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_me_requires_bearer_token() -> None:
    response = client.get("/me")

    assert response.status_code == 401
    assert response.json()["detail"] == "Missing bearer token."


def test_me_rejects_invalid_dev_token() -> None:
    response = client.get("/me", headers={"Authorization": "Bearer wrong-token"})

    assert response.status_code == 401
    assert response.json()["detail"] == "Invalid bearer token."


def test_me_returns_current_user_with_valid_dev_token() -> None:
    response = client.get("/me", headers=AUTH_HEADERS)

    assert response.status_code == 200
    body = response.json()
    assert body == {
        "user_id": "test-clerk-user-1",
        "auth_provider": "dev",
    }


def test_public_smart_fit_still_works_without_login() -> None:
    response = client.post(
        "/analysis/smart",
        json={
            "resume_text": "PROJECTS\nBuilt Python FastAPI services with Docker, SQL, Git, and REST APIs.",
            "job_description": "Required Qualifications\nBuild Python REST APIs with SQL, Docker, Git, and FastAPI.",
        },
    )

    assert response.status_code == 200
    assert response.json()["fit_summary"]["score"] >= 1


def test_create_job_posting_requires_admin_api_key() -> None:
    response = client.post("/job-postings", json=_sample_job_posting_payload())

    assert response.status_code == 401
    assert response.json()["detail"] == "Invalid or missing admin API key."


def test_create_job_posting_extracts_and_persists_skills() -> None:
    response = client.post(
        "/job-postings",
        json=_sample_job_posting_payload(),
        headers=ADMIN_HEADERS,
    )

    assert response.status_code == 201
    created_posting = response.json()
    assert created_posting["company"] == "Comcast"
    assert created_posting["extracted_skills"] == [
        "Agile",
        "Docker",
        "Git",
        "Python",
        "REST APIs",
        "SQL",
    ]

    list_response = client.get("/job-postings")
    assert list_response.status_code == 200
    assert len(list_response.json()) == 1


def test_csv_import_requires_admin_api_key() -> None:
    csv_content = (
        "company,title,location,role_category,experience_level,description\n"
        "Lockheed Martin,Software Engineer Associate,King of Prussia PA,Systems/SWE,Entry-Level,Python Linux Git Agile testing\n"
    )

    response = client.post(
        "/job-postings/import-csv",
        files={"file": ("jobs.csv", csv_content.encode("utf-8"), "text/csv")},
    )

    assert response.status_code == 401
    assert response.json()["detail"] == "Invalid or missing admin API key."


def test_csv_import_creates_multiple_postings() -> None:
    csv_content = (
        "company,title,location,role_category,experience_level,description\n"
        "Lockheed Martin,Software Engineer Associate,King of Prussia PA,Systems/SWE,Entry-Level,Python Linux Git Agile testing\n"
        "Leidos,Cloud Automation Engineer,Remote,Cloud/SWE,Entry-Level,Python AWS Docker CI/CD automation scripting\n"
    )

    response = client.post(
        "/job-postings/import-csv",
        files={"file": ("jobs.csv", csv_content.encode("utf-8"), "text/csv")},
        headers=ADMIN_HEADERS,
    )

    assert response.status_code == 201
    body = response.json()
    assert body["imported_count"] == 2
    assert body["failed_count"] == 0
    assert len(body["created_postings"]) == 2

    top_skills_response = client.get("/skills/top")
    assert top_skills_response.status_code == 200
    top_skills = top_skills_response.json()
    assert top_skills["Python"] == 2
    assert top_skills["Docker"] == 1
    assert top_skills["AWS"] == 1


def test_resume_file_extracts_plain_text_without_saving() -> None:
    response = client.post(
        "/analysis/resume-file/extract",
        files={
            "file": (
                "resume.txt",
                b"Python FastAPI Docker SQL\nBuilt backend APIs.",
                "text/plain",
            )
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["filename"] == "resume.txt"
    assert "Python FastAPI Docker SQL" in body["text"]
    assert body["character_count"] == len(body["text"])
    assert "not saved" in " ".join(body["warnings"])

    saved_postings_response = client.get("/job-postings")
    assert saved_postings_response.status_code == 200
    assert saved_postings_response.json() == []


def test_resume_file_extracts_docx_text() -> None:
    response = client.post(
        "/analysis/resume-file/extract",
        files={
            "file": (
                "resume.docx",
                _sample_docx_resume_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["filename"] == "resume.docx"
    assert "Python FastAPI services" in body["text"]
    assert "Docker" in body["text"]


def test_resume_file_rejects_unsupported_file_type() -> None:
    response = client.post(
        "/analysis/resume-file/extract",
        files={"file": ("resume.rtf", b"Python and SQL", "application/rtf")},
    )

    assert response.status_code == 400
    assert ".txt, .md, .pdf, and .docx" in response.json()["detail"]


def test_resume_analysis_compares_resume_against_target_role_category() -> None:
    client.post(
        "/job-postings",
        json={
            "company": "Comcast",
            "title": "Backend Software Engineer",
            "location": "Philadelphia, PA",
            "role_category": "Backend SWE",
            "experience_level": "Entry-Level",
            "description": "Backend role using Python, SQL, Docker, and REST APIs.",
        },
        headers=ADMIN_HEADERS,
    )
    client.post(
        "/job-postings",
        json={
            "company": "UHS",
            "title": "Systems Engineer Intern",
            "location": "King of Prussia, PA",
            "role_category": "Systems/Cloud",
            "experience_level": "Internship",
            "description": "Systems role using Azure, Windows Server, scripting, and troubleshooting.",
        },
        headers=ADMIN_HEADERS,
    )

    response = client.post(
        "/resume/analyze",
        json={
            "resume_text": "Python, SQL, Git, Agile, and backend development projects.",
            "target_role_category": "Backend SWE",
        },
    )

    assert response.status_code == 200
    analysis = response.json()
    assert analysis["postings_analyzed"] == 1
    assert analysis["target_role_category"] == "Backend SWE"
    assert analysis["matched_skills"] == ["Python", "SQL"]
    assert analysis["missing_skills"] == ["Docker", "REST APIs"]
    assert analysis["match_percentage"] == 50.0
    assert analysis["learning_priorities"] == ["Docker", "REST APIs"]


def test_resume_analysis_returns_error_when_target_role_has_no_postings() -> None:
    response = client.post(
        "/resume/analyze",
        json={
            "resume_text": "Python, SQL, and Docker",
            "target_role_category": "AI Engineer",
        },
    )

    assert response.status_code == 400
    assert response.json()["detail"] == "No saved job postings found for that target role category."


def test_resume_analysis_rejects_overly_long_resume_text() -> None:
    response = client.post(
        "/resume/analyze",
        json={"resume_text": "x" * 10_001},
    )

    assert response.status_code == 422


def test_custom_analysis_compares_resume_against_pasted_job_descriptions_without_saving() -> None:
    response = client.post(
        "/analysis/custom",
        json={
            "resume_text": "Python, SQL, Git, Agile, and backend development projects.",
            "job_descriptions": [
                "Backend role using Python, SQL, Docker, and REST APIs.",
            ],
        },
    )

    assert response.status_code == 200
    analysis = response.json()
    assert analysis["postings_analyzed"] == 1
    assert analysis["target_role_category"] is None
    assert analysis["matched_skills"] == ["Python", "SQL"]
    assert analysis["missing_skills"] == ["Docker", "REST APIs"]
    assert analysis["match_percentage"] == 50.0
    assert analysis["learning_priorities"] == ["Docker", "REST APIs"]

    saved_postings_response = client.get("/job-postings")
    assert saved_postings_response.status_code == 200
    assert saved_postings_response.json() == []


def test_custom_analysis_returns_error_when_no_target_skills_are_found() -> None:
    response = client.post(
        "/analysis/custom",
        json={
            "resume_text": "Python and SQL",
            "job_descriptions": ["Friendly team with strong communication and collaboration."],
        },
    )

    assert response.status_code == 400
    assert response.json()["detail"] == "No recognizable target skills found in the provided job descriptions."


def test_smart_fit_batch_endpoint_ranks_jobs_independently_without_saving() -> None:
    response = client.post(
        "/analysis/smart/batch",
        json={
            "resume_text": "PROJECTS\nBuilt Python FastAPI services with Docker, SQL, Git, and REST APIs.",
            "job_descriptions": [
                {
                    "title": "Cloud Stretch Role",
                    "job_description": "Required Qualifications\nBuild AWS Kubernetes systems with CI/CD and Terraform.",
                },
                {
                    "title": "Backend Match Role",
                    "job_description": "Required Qualifications\nBuild Python REST APIs with SQL, Docker, Git, and FastAPI.",
                },
            ],
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["analyzed_count"] == 2
    assert [result["rank"] for result in body["results"]] == [1, 2]
    assert body["best_job"]["title"] == "Backend Match Role"
    assert body["results"][0]["analysis"]["fit_summary"]["score"] >= body["results"][1]["analysis"]["fit_summary"]["score"]

    saved_postings_response = client.get("/job-postings")
    assert saved_postings_response.status_code == 200
    assert saved_postings_response.json() == []


def test_smart_fit_batch_endpoint_identifies_bad_job_input() -> None:
    response = client.post(
        "/analysis/smart/batch",
        json={
            "resume_text": "Python and SQL project work.",
            "job_descriptions": [
                {"title": "Bad Input", "job_description": "Friendly team culture only."},
            ],
        },
    )

    assert response.status_code == 400
    assert response.json()["detail"].startswith("Job 1:")


def test_model_assisted_status_reports_disabled_without_secrets(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.delenv("AI_ANALYSIS_ENABLED", raising=False)
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    monkeypatch.delenv("OPENAI_MODEL", raising=False)

    response = client.get("/analysis/model-status")

    assert response.status_code == 200
    body = response.json()
    assert body["enabled"] is False
    assert body["status"] == "not_configured"
    assert "OPENAI_API_KEY" in body["required_backend_settings"]
    assert "Provider keys must stay in backend environment variables only." in body["safety_notes"]


def test_model_assisted_status_never_exposes_backend_secret(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("AI_ANALYSIS_ENABLED", "true")
    monkeypatch.setenv("OPENAI_API_KEY", "super-secret-test-key")
    monkeypatch.setenv("OPENAI_MODEL", "gpt-test")

    response = client.get("/analysis/model-status")

    assert response.status_code == 200
    body = response.json()
    assert body["enabled"] is True
    assert body["status"] == "configured"
    assert "super-secret-test-key" not in str(body)


def _sample_saved_job_payload() -> dict[str, str]:
    return {
        "source": "greenhouse",
        "source_job_id": "job-123",
        "company": "Example Health",
        "title": "Software Engineer",
        "location": "Philadelphia, PA",
        "description": "Build Python REST APIs using SQL, Docker, Git, and Agile practices.",
        "apply_url": "https://example.com/jobs/123",
    }


def test_saved_jobs_require_authentication() -> None:
    list_response = client.get("/saved-jobs")
    create_response = client.post(
        "/saved-jobs",
        json=_sample_saved_job_payload(),
    )

    assert list_response.status_code == 401
    assert create_response.status_code == 401


def test_authenticated_user_can_create_and_list_saved_jobs() -> None:
    create_response = client.post(
        "/saved-jobs",
        json=_sample_saved_job_payload(),
        headers=AUTH_HEADERS,
    )

    assert create_response.status_code == 201
    created_job = create_response.json()
    assert created_job["company"] == "Example Health"
    assert created_job["source_job_id"] == "job-123"
    assert created_job["extracted_skills"] == [
        "Agile",
        "Docker",
        "Git",
        "Python",
        "REST APIs",
        "SQL",
    ]
    assert "user_id" not in created_job

    list_response = client.get("/saved-jobs", headers=AUTH_HEADERS)

    assert list_response.status_code == 200
    assert list_response.json() == [created_job]


def test_saved_jobs_are_isolated_by_authenticated_user(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    first_user_response = client.post(
        "/saved-jobs",
        json=_sample_saved_job_payload(),
        headers=AUTH_HEADERS,
    )
    first_user_job_id = first_user_response.json()["id"]

    monkeypatch.setenv("AUTH_DEV_USER_ID", "test-clerk-user-2")

    second_user_payload = {
        **_sample_saved_job_payload(),
        "source_job_id": "job-456",
        "company": "Second User Company",
        "title": "Data Engineer",
    }
    second_user_response = client.post(
        "/saved-jobs",
        json=second_user_payload,
        headers=AUTH_HEADERS,
    )

    assert second_user_response.status_code == 201

    second_user_list = client.get("/saved-jobs", headers=AUTH_HEADERS)
    assert second_user_list.status_code == 200
    assert len(second_user_list.json()) == 1
    assert second_user_list.json()[0]["company"] == "Second User Company"

    hidden_job_response = client.get(
        f"/saved-jobs/{first_user_job_id}",
        headers=AUTH_HEADERS,
    )
    hidden_delete_response = client.delete(
        f"/saved-jobs/{first_user_job_id}",
        headers=AUTH_HEADERS,
    )

    assert hidden_job_response.status_code == 404
    assert hidden_delete_response.status_code == 404


def test_authenticated_user_can_delete_own_saved_job() -> None:
    create_response = client.post(
        "/saved-jobs",
        json=_sample_saved_job_payload(),
        headers=AUTH_HEADERS,
    )
    saved_job_id = create_response.json()["id"]

    delete_response = client.delete(
        f"/saved-jobs/{saved_job_id}",
        headers=AUTH_HEADERS,
    )
    get_response = client.get(
        f"/saved-jobs/{saved_job_id}",
        headers=AUTH_HEADERS,
    )

    assert delete_response.status_code == 200
    assert delete_response.json() == {"status": "deleted"}
    assert get_response.status_code == 404


def test_saving_same_external_job_twice_is_idempotent() -> None:
    first_response = client.post(
        "/saved-jobs",
        json=_sample_saved_job_payload(),
        headers=AUTH_HEADERS,
    )
    second_response = client.post(
        "/saved-jobs",
        json=_sample_saved_job_payload(),
        headers=AUTH_HEADERS,
    )

    assert first_response.status_code == 201
    assert second_response.status_code == 201
    assert first_response.json()["id"] == second_response.json()["id"]

    list_response = client.get("/saved-jobs", headers=AUTH_HEADERS)

    assert list_response.status_code == 200
    assert len(list_response.json()) == 1
