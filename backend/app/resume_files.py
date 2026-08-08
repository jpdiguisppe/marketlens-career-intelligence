\
"""Resume file text extraction helpers with bounded parser resource use."""

from __future__ import annotations

import os
from io import BytesIO
from pathlib import PurePosixPath
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader, filters as pypdf_filters
from pypdf.errors import PdfReadError

SUPPORTED_RESUME_UPLOAD_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}

MAX_DOCX_ARCHIVE_ENTRIES = 512
MAX_DOCX_EXPANDED_BYTES = 8_000_000
MAX_DOCX_ENTRY_BYTES = 4_000_000
MAX_DOCX_COMPRESSION_RATIO = 100
MAX_DOCX_PARAGRAPHS = 5_000
MAX_DOCX_TABLE_CELLS = 10_000
MAX_PDF_PAGES = 20
MAX_PDF_PAGE_CONTENT_STREAM_BYTES = 8_000_000
MAX_PARSER_EXTRACTED_TEXT_CHARACTERS = 50_000

# pypdf documents its decompression controls as security knobs. MarketLens
# processes small resumes, so lower the library-wide PDF decode ceilings
# from the much larger defaults before any uploaded page is decoded.
for _filter_limit_name in (
    "ZLIB_MAX_OUTPUT_LENGTH",
    "LZW_MAX_OUTPUT_LENGTH",
    "RUN_LENGTH_MAX_OUTPUT_LENGTH",
    "JBIG2_MAX_OUTPUT_LENGTH",
    "MAX_DECLARED_STREAM_LENGTH",
    "MAX_ARRAY_BASED_STREAM_OUTPUT_LENGTH",
    "FLATE_MAX_BUFFER_SIZE",
):
    if hasattr(pypdf_filters, _filter_limit_name):
        setattr(pypdf_filters, _filter_limit_name, MAX_PDF_PAGE_CONTENT_STREAM_BYTES)
if hasattr(pypdf_filters, "JBIG2DEC_BINARY"):
    pypdf_filters.JBIG2DEC_BINARY = None


class ResumeFileExtractionError(ValueError):
    """Raised when uploaded resume text cannot be extracted safely."""


def _normalize_extracted_text(text: str) -> str:
    return "\n".join(line.rstrip() for line in text.splitlines()).strip()


def _extract_text_file(contents: bytes) -> str:
    try:
        return contents.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ResumeFileExtractionError("Resume text file must be UTF-8 encoded.") from exc


def _validate_docx_archive(contents: bytes) -> None:
    try:
        with ZipFile(BytesIO(contents)) as archive:
            entries = [entry for entry in archive.infolist() if not entry.is_dir()]
            if len(entries) > MAX_DOCX_ARCHIVE_ENTRIES:
                raise ResumeFileExtractionError(
                    "DOCX resume contains too many internal files to process safely."
                )

            names = {entry.filename for entry in entries}
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise ResumeFileExtractionError("DOCX resume structure is invalid.")

            expanded_bytes = 0
            for entry in entries:
                path = PurePosixPath(entry.filename)
                if entry.filename.startswith("/") or ".." in path.parts:
                    raise ResumeFileExtractionError("DOCX resume contains an unsafe archive path.")
                if entry.flag_bits & 0x1:
                    raise ResumeFileExtractionError("Encrypted DOCX resumes are not supported.")
                if entry.file_size > MAX_DOCX_ENTRY_BYTES:
                    raise ResumeFileExtractionError(
                        "DOCX resume contains an oversized internal file."
                    )

                expanded_bytes += entry.file_size
                if expanded_bytes > MAX_DOCX_EXPANDED_BYTES:
                    raise ResumeFileExtractionError(
                        "DOCX resume expands beyond the safe processing limit."
                    )

                if entry.file_size >= 64 * 1024:
                    ratio = entry.file_size / max(entry.compress_size, 1)
                    if ratio > MAX_DOCX_COMPRESSION_RATIO:
                        raise ResumeFileExtractionError(
                            "DOCX resume has a suspicious compression ratio."
                        )

            for entry in entries:
                if not entry.filename.lower().endswith((".xml", ".rels")):
                    continue
                xml_bytes = archive.read(entry)
                upper_xml = xml_bytes.upper()
                if b"<!DOCTYPE" in upper_xml or b"<!ENTITY" in upper_xml:
                    raise ResumeFileExtractionError(
                        "DOCX resume contains unsupported XML entity declarations."
                    )
    except ResumeFileExtractionError:
        raise
    except (BadZipFile, OSError, ValueError) as exc:
        raise ResumeFileExtractionError("Could not read this DOCX resume file.") from exc


def _extract_docx_text(contents: bytes) -> str:
    if not contents.startswith(b"PK"):
        raise ResumeFileExtractionError("DOCX resume structure is invalid.")
    _validate_docx_archive(contents)

    try:
        document = Document(BytesIO(contents))
    except Exception as exc:  # python-docx exposes several zip/XML parser exceptions.
        raise ResumeFileExtractionError("Could not read text from this DOCX resume file.") from exc

    if len(document.paragraphs) > MAX_DOCX_PARAGRAPHS:
        raise ResumeFileExtractionError("DOCX resume contains too many paragraphs to process safely.")
    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    table_cells: list[str] = []
    cell_count = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_count += 1
                if cell_count > MAX_DOCX_TABLE_CELLS:
                    raise ResumeFileExtractionError(
                        "DOCX resume contains too many table cells to process safely."
                    )
                if cell.text.strip():
                    table_cells.append(cell.text)

    return "\n".join(paragraphs + table_cells)


def _extract_pdf_text(contents: bytes) -> str:
    if not contents.lstrip().startswith(b"%PDF-"):
        raise ResumeFileExtractionError("PDF resume structure is invalid.")

    try:
        reader = PdfReader(BytesIO(contents))
        if reader.is_encrypted:
            raise ResumeFileExtractionError("Encrypted PDF resumes are not supported.")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ResumeFileExtractionError(
                f"PDF resume has too many pages. Maximum supported pages: {MAX_PDF_PAGES}."
            )

        page_text: list[str] = []
        extracted_characters = 0
        for page in reader.pages:
            page_contents = page.get_contents()
            if page_contents is not None:
                decoded_stream = page_contents.get_data()
                if len(decoded_stream) > MAX_PDF_PAGE_CONTENT_STREAM_BYTES:
                    raise ResumeFileExtractionError(
                        "PDF resume contains an oversized page content stream."
                    )

            extracted_page_text = page.extract_text() or ""
            extracted_characters += len(extracted_page_text)
            if extracted_characters > MAX_PARSER_EXTRACTED_TEXT_CHARACTERS:
                raise ResumeFileExtractionError(
                    "PDF resume contains too much extracted text to process safely."
                )
            if extracted_page_text.strip():
                page_text.append(extracted_page_text)

        return "\n".join(page_text)
    except ResumeFileExtractionError:
        raise
    except (PdfReadError, ValueError, OSError) as exc:
        raise ResumeFileExtractionError("Could not read text from this PDF resume file.") from exc
    except Exception as exc:
        raise ResumeFileExtractionError("PDF resume exceeded safe parser limits.") from exc


def extract_resume_text_from_upload(filename: str, contents: bytes) -> tuple[str, list[str]]:
    """Extract readable text and warnings from a supported resume upload."""

    extension = os.path.splitext(filename.lower())[1]
    warnings = [
        "Uploaded resume text is returned for this request and is not saved to the shared database."
    ]

    if extension not in SUPPORTED_RESUME_UPLOAD_EXTENSIONS:
        raise ResumeFileExtractionError(
            "Resume upload supports .txt, .md, .pdf, and .docx files."
        )

    if extension in {".txt", ".md"}:
        raw_text = _extract_text_file(contents)
    elif extension == ".docx":
        raw_text = _extract_docx_text(contents)
    elif extension == ".pdf":
        raw_text = _extract_pdf_text(contents)
        warnings.append(
            "PDF extraction depends on whether the PDF contains selectable text; scanned/image-only PDFs may not work."
        )
    else:
        raise ResumeFileExtractionError(
            "Resume upload supports .txt, .md, .pdf, and .docx files."
        )

    text = _normalize_extracted_text(raw_text)
    if not text:
        raise ResumeFileExtractionError(
            "Uploaded resume file did not contain readable text. Try exporting it as DOCX or plain text."
        )

    return text, warnings
